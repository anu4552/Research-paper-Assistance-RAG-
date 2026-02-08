#  app.py – Backend Ollama Summary + DOCX + PDF

from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
import os, uuid, re, requests, traceback, base64
from jinja2 import Template
from markupsafe import Markup
from weasyprint import HTML
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "uploads"
TEMPLATE_PATH = "preview.html"
PROMPT_PATH = "prompt_template.txt"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Load HTML template on startup
with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
    html_template = Template(f.read())

# Load prompt template
with open(PROMPT_PATH, "r") as f:
    prompt_template = f.read()

def generate_summary(data):
    prompt = prompt_template.format(**data)
    response = requests.post("http://localhost:11434/api/generate", json={
        "model": "deepseek-r1:1.5b",
        "prompt": prompt,
        "temperature": 0, 
        "stream": False
    })
    if response.status_code != 200:
        raise Exception(f"Ollama error: {response.status_code}: {response.text}")

    summary = response.json().get("response", "").strip()
    summary = re.sub(r"<think>.*?</think>", "", summary, flags=re.DOTALL)
    summary = re.sub(r"\**Summary\**:?", "", summary, flags=re.IGNORECASE).strip()
    return summary

def save_base64_image(base64_data):
    header, encoded = base64_data.split(",", 1)
    ext = header.split("/")[1].split(";")[0]
    filename = f"{uuid.uuid4().hex}.{ext}"
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    with open(filepath, "wb") as f:
        f.write(base64.b64decode(encoded))
    return filepath

@app.route("/generate-preview", methods=["POST"])
def generate_preview():
    try:
        data = request.form.to_dict()
        photo_url = data.get("PHOTO_URL")

        if not photo_url:
            photo = request.files.get("photo")
            if photo:
                photo_ext = photo.filename.split(".")[-1]
                photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
                photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
                photo.save(photo_path)
                photo_url = f"/uploads/{photo_filename}"
            else:
                photo_url = "https://via.placeholder.com/120"

        data["PHOTO_URL"] = photo_url
        data["SUMMARY"] = generate_summary(data)

        context = {
            "NAME": data.get("name", ""),
            "ADDRESS": data.get("address", ""),
            "EMAIL": data.get("email", ""),
            "NUMBER": data.get("phone", ""),
            "ROLE": data.get("role", ""),
            "EXPERIENCE": Markup(data.get("experience", "")),
            "EDUCATION": Markup(data.get("education", "")),
            "PROFESSIONAL": data.get("prof_education", ""),
            "SKILLS": Markup(data.get("skills", "")),
            "PROJECTS": Markup(data.get("projects", "")),
            "AWARDS": Markup(data.get("awards", "")),
            "START": data.get("start", ""),
            "END": data.get("end", ""),
            "SUMMARY": data.get("SUMMARY"),
            "PHOTO_URL": photo_url
        }

        html_out = html_template.render(**context)
        return jsonify({"html": html_out, "summary": data.get("SUMMARY")})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/download-pdf", methods=["POST"])
def download_pdf():
    try:
        html = request.form.get("html", "")
        if not html:
            return jsonify({"error": "No HTML provided"}), 400

        pdf_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.pdf")
        HTML(string=html, base_url=os.getcwd()).write_pdf(pdf_path)

        return send_file(pdf_path, as_attachment=True, download_name="resume.pdf")

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    # try:
    #     data = request.form.to_dict()
    #     summary = data.get("SUMMARY", "")

    #     photo_url = data.get("PHOTO_URL")
    #     if not photo_url:
    #         photo = request.files.get("photo")
    #         if photo:
    #             photo_ext = photo.filename.split(".")[-1]
    #             photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
    #             photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
    #             photo.save(photo_path)
    #             photo_url = f"file://{os.path.abspath(photo_path)}"
    #         else:
    #             photo_url = "https://via.placeholder.com/120"
    #     data["PHOTO_URL"] = photo_url

    #     context = {
    #         "NAME": data.get("name", ""),
    #         "ADDRESS": data.get("address", ""),
    #         "EMAIL": data.get("email", ""),
    #         "NUMBER": data.get("phone", ""),
    #         "ROLE": data.get("role", ""),
    #         "EXPERIENCE": Markup(data.get("experience", "")),
    #         "EDUCATION": Markup(data.get("education", "")),
    #         "PROFESSIONAL": data.get("prof_education", ""),
    #         "SKILLS": Markup(data.get("skills", "")),
    #         "PROJECTS": Markup(data.get("projects", "")),
    #         "AWARDS": Markup(data.get("awards", "")),
    #         "START": data.get("start", ""),
    #         "END": data.get("end", ""),
    #         "SUMMARY": summary,
    #         "PHOTO_URL": photo_url
    #     }

    #     final_html = html_template.render(**context)
    #     pdf_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.pdf")
    #     HTML(string=final_html, base_url=os.getcwd()).write_pdf(pdf_path)
    #     return send_file(pdf_path, as_attachment=True, download_name="resume.pdf")

    # except Exception as e:
    #     traceback.print_exc()
    #     return jsonify({"error": str(e)}), 500

@app.route("/download-docx", methods=["POST"])
def download_docx():
    try:
        data = request.form.to_dict()
        summary = data.get("SUMMARY", "")

        doc = Document()
        doc.add_heading(data.get("name", ""), 0)
        doc.add_paragraph(f"Email: {data.get('email', '')}")
        doc.add_paragraph(f"Phone: {data.get('phone', '')}")
        doc.add_paragraph(f"Address: {data.get('address', '')}")
        doc.add_paragraph(f"Role: {data.get('role', '')}")
        doc.add_paragraph("Summary:")
        doc.add_paragraph(summary)
        doc.add_paragraph("Experience:")
        doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("experience", "")))
        doc.add_paragraph("Education:")
        doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("education", "")))
        doc.add_paragraph("Skills:")
        doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("skills", "")))
        doc.add_paragraph("Projects:")
        doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("projects", "")))
        doc.add_paragraph("Awards:")
        doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("awards", "")))

        photo_url = data.get("PHOTO_URL", "")
        if photo_url.startswith("data:image"):
            image_path = save_base64_image(photo_url)
            doc.add_picture(image_path, width=Inches(1.5))

        output_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.docx")
        doc.save(output_path)
        return send_file(output_path, as_attachment=True, download_name="resume.docx")

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/uploads/<filename>")
def uploaded_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

if __name__ == "__main__":
    app.run(debug=True)



# # ✅ app.py – Backend for CV Generator with HTML Templating + Ollama Summary + DOCX + PDF

# from flask import Flask, request, jsonify, send_file, send_from_directory
# from flask_cors import CORS
# import os, uuid, re, requests, traceback, base64
# from jinja2 import Template
# from markupsafe import Markup
# from weasyprint import HTML
# from docx import Document
# from docx.shared import Inches

# app = Flask(__name__)
# CORS(app)

# UPLOAD_FOLDER = "uploads"
# TEMPLATE_PATH = "preview.html"
# PROMPT_PATH = "prompt_template.txt"

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# # Load HTML template on startup
# with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
#     html_template = Template(f.read())

# # Load prompt template
# with open(PROMPT_PATH, "r") as f:
#     prompt_template = f.read()

# def generate_summary(data):
#     prompt = prompt_template.format(**data)
#     response = requests.post("http://localhost:11434/api/generate", json={
#         "model": "deepseek-r1:1.5b",
#         "prompt": prompt,
#         "stream": False
#     })
#     if response.status_code != 200:
#         raise Exception(f"Ollama error: {response.status_code}: {response.text}")

#     summary = response.json().get("response", "").strip()
#     summary = re.sub(r"<think>.*?</think>", "", summary, flags=re.DOTALL)
#     return summary

# def save_base64_image(base64_data):
#     header, encoded = base64_data.split(",", 1)
#     ext = header.split("/")[1].split(";")[0]
#     filename = f"{uuid.uuid4().hex}.{ext}"
#     filepath = os.path.join(UPLOAD_FOLDER, filename)
#     with open(filepath, "wb") as f:
#         f.write(base64.b64decode(encoded))
#     return filepath

# @app.route("/generate-preview", methods=["POST"])
# def generate_preview():
#     try:
#         data = request.form.to_dict()
#         photo_url = data.get("PHOTO_URL")

#         if not photo_url:
#             photo = request.files.get("photo")
#             if photo:
#                 photo_ext = photo.filename.split(".")[-1]
#                 photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
#                 photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
#                 photo.save(photo_path)
#                 photo_url = f"/uploads/{photo_filename}"
#             else:
#                 photo_url = "https://via.placeholder.com/120"

#         data["PHOTO_URL"] = photo_url
#         data["SUMMARY"] = generate_summary(data)

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": Markup(data.get("experience", "")),
#             "EDUCATION": Markup(data.get("education", "")),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": Markup(data.get("skills", "")),
#             "PROJECTS": Markup(data.get("projects", "")),
#             "AWARDS": Markup(data.get("awards", "")),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": data.get("SUMMARY"),
#             "PHOTO_URL": photo_url
#         }

#         html_out = html_template.render(**context)
#         return jsonify({"html": html_out})

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/download-pdf", methods=["POST"])
# def download_pdf():
#     try:
#         data = request.form.to_dict()
#         data["SUMMARY"] = generate_summary(data)

#         photo_url = data.get("PHOTO_URL")
#         if not photo_url:
#             photo = request.files.get("photo")
#             if photo:
#                 photo_ext = photo.filename.split(".")[-1]
#                 photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
#                 photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
#                 photo.save(photo_path)
#                 photo_url = f"file://{os.path.abspath(photo_path)}"
#             else:
#                 photo_url = "https://via.placeholder.com/120"
#         data["PHOTO_URL"] = photo_url

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": Markup(data.get("experience", "")),
#             "EDUCATION": Markup(data.get("education", "")),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": Markup(data.get("skills", "")),
#             "PROJECTS": Markup(data.get("projects", "")),
#             "AWARDS": Markup(data.get("awards", "")),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": data.get("SUMMARY"),
#             "PHOTO_URL": data.get("PHOTO_URL")
#         }

#         final_html = html_template.render(**context)
#         pdf_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.pdf")
#         HTML(string=final_html, base_url=os.getcwd()).write_pdf(pdf_path)
#         return send_file(pdf_path, as_attachment=True, download_name="resume.pdf")

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/download-docx", methods=["POST"])
# def download_docx():
#     try:
#         data = request.form.to_dict()
#         data["SUMMARY"] = generate_summary(data)

#         doc = Document()
#         doc.add_heading(data.get("name", ""), 0)
#         doc.add_paragraph(f"Email: {data.get('email', '')}")
#         doc.add_paragraph(f"Phone: {data.get('phone', '')}")
#         doc.add_paragraph(f"Address: {data.get('address', '')}")
#         doc.add_paragraph(f"Role: {data.get('role', '')}")
#         doc.add_paragraph("Summary:")
#         doc.add_paragraph(data.get("SUMMARY", ""))
#         doc.add_paragraph("Experience:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("experience", "")))
#         doc.add_paragraph("Education:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("education", "")))
#         doc.add_paragraph("Skills:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("skills", "")))
#         doc.add_paragraph("Projects:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("projects", "")))
#         doc.add_paragraph("Awards:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("awards", "")))

#         photo_url = data.get("PHOTO_URL", "")
#         if photo_url.startswith("data:image"):
#             image_path = save_base64_image(photo_url)
#             doc.add_picture(image_path, width=Inches(1.5))

#         output_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.docx")
#         doc.save(output_path)
#         return send_file(output_path, as_attachment=True, download_name="resume.docx")

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/uploads/<filename>")
# def uploaded_file(filename):
#     return send_from_directory(UPLOAD_FOLDER, filename)

# if __name__ == "__main__":
#     app.run(debug=True)






# # ✅ app.py – Backend for CV Generator with HTML Templating + Ollama Summary + DOCX + PDF

# from flask import Flask, request, jsonify, send_file, send_from_directory
# from flask_cors import CORS
# import os, uuid, re, requests, traceback
# from jinja2 import Template
# from markupsafe import Markup
# from weasyprint import HTML
# from docx import Document
# from docx.shared import Inches

# app = Flask(__name__)
# CORS(app)

# UPLOAD_FOLDER = "uploads"
# TEMPLATE_PATH = "preview.html"
# PROMPT_PATH = "prompt_template.txt"

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# # Load HTML template on startup
# with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
#     html_template = Template(f.read())

# # Load prompt template
# with open(PROMPT_PATH, "r") as f:
#     prompt_template = f.read()

# def generate_summary(data):
#     prompt = prompt_template.format(**data)
#     response = requests.post("http://localhost:11434/api/generate", json={
#         "model": "deepseek-r1:1.5b",
#         "prompt": prompt,
#         "stream": False
#     })
#     if response.status_code != 200:
#         raise Exception(f"Ollama error: {response.status_code}: {response.text}")

#     summary = response.json().get("response", "").strip()
#     summary = re.sub(r"<think>.*?</think>", "", summary, flags=re.DOTALL)
#     return summary

# @app.route("/generate-preview", methods=["POST"])
# def generate_preview():
#     try:
#         data = request.form.to_dict()
#         photo_url = data.get("PHOTO_URL")

#         if not photo_url:
#             photo = request.files.get("photo")
#             if photo:
#                 photo_ext = photo.filename.split(".")[-1]
#                 photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
#                 photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
#                 photo.save(photo_path)
#                 photo_url = f"/uploads/{photo_filename}"
#             else:
#                 photo_url = "https://via.placeholder.com/120"

#         data["PHOTO_URL"] = photo_url
#         data["SUMMARY"] = generate_summary(data)

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": Markup(data.get("experience", "")),
#             "EDUCATION": Markup(data.get("education", "")),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": Markup(data.get("skills", "")),
#             "PROJECTS": Markup(data.get("projects", "")),
#             "AWARDS": Markup(data.get("awards", "")),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": data.get("SUMMARY"),
#             "PHOTO_URL": photo_url
#         }

#         html_out = html_template.render(**context)
#         return jsonify({"html": html_out})

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/download-pdf", methods=["POST"])
# def download_pdf():
#     try:
#         data = request.form.to_dict()
#         data["SUMMARY"] = generate_summary(data)

#         photo_url = data.get("PHOTO_URL")
#         if not photo_url:
#             photo = request.files.get("photo")
#             if photo:
#                 photo_ext = photo.filename.split(".")[-1]
#                 photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
#                 photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
#                 photo.save(photo_path)
#                 photo_url = f"file://{os.path.abspath(photo_path)}"
#             else:
#                 photo_url = "https://via.placeholder.com/120"
#         data["PHOTO_URL"] = photo_url

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": Markup(data.get("experience", "")),
#             "EDUCATION": Markup(data.get("education", "")),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": Markup(data.get("skills", "")),
#             "PROJECTS": Markup(data.get("projects", "")),
#             "AWARDS": Markup(data.get("awards", "")),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": data.get("SUMMARY"),
#             "PHOTO_URL": data.get("PHOTO_URL")
#         }

#         final_html = html_template.render(**context)
#         pdf_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.pdf")
#         HTML(string=final_html, base_url=os.getcwd()).write_pdf(pdf_path)
#         return send_file(pdf_path, as_attachment=True, download_name="resume.pdf")

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/download-docx", methods=["POST"])
# def download_docx():
#     try:
#         data = request.form.to_dict()
#         data["SUMMARY"] = generate_summary(data)

#         doc = Document()
#         doc.add_heading(data.get("name", ""), 0)
#         doc.add_paragraph(f"Email: {data.get('email', '')}")
#         doc.add_paragraph(f"Phone: {data.get('phone', '')}")
#         doc.add_paragraph(f"Address: {data.get('address', '')}")
#         doc.add_paragraph(f"Role: {data.get('role', '')}")
#         doc.add_paragraph("Summary:")
#         doc.add_paragraph(data.get("SUMMARY", ""))
#         doc.add_paragraph("Experience:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("experience", "")))
#         doc.add_paragraph("Education:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("education", "")))
#         doc.add_paragraph("Skills:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("skills", "")))
#         doc.add_paragraph("Projects:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("projects", "")))
#         doc.add_paragraph("Awards:")
#         doc.add_paragraph(re.sub('<[^<]+?>', '', data.get("awards", "")))

#         output_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.docx")
#         doc.save(output_path)
#         return send_file(output_path, as_attachment=True, download_name="resume.docx")

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/uploads/<filename>")
# def uploaded_file(filename):
#     return send_from_directory(UPLOAD_FOLDER, filename)

# if __name__ == "__main__":
#     app.run(debug=True)







# # ✅ app.py – Backend for CV Generator with HTML Templating + Ollama Summary

# from flask import Flask, request, jsonify, send_file, send_from_directory
# from flask_cors import CORS
# import os, uuid, re, requests, traceback
# from jinja2 import Template
# from markupsafe import Markup
# from werkzeug.utils import secure_filename
# from weasyprint import HTML

# app = Flask(__name__)
# CORS(app)

# UPLOAD_FOLDER = os.path.join(os.path.abspath(os.path.dirname(__file__)), "uploads")
# TEMPLATE_PATH = "preview.html"
# PROMPT_PATH = "prompt_template.txt"

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# # Load HTML template on startup
# with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
#     html_template = Template(f.read())

# # Load prompt template
# with open(PROMPT_PATH, "r") as f:
#     prompt_template = f.read()

# def generate_summary(data):
#     prompt = prompt_template.format(**data)
#     print("🧠 Prompt sent to Ollama:\n", prompt)
#     response = requests.post("http://localhost:11434/api/generate", json={
#         "model": "deepseek-r1:1.5b",
#         "prompt": prompt,
#         "stream": False
#     })
#     if response.status_code != 200:
#         raise Exception(f"Ollama error: {response.status_code}: {response.text}")
#     summary = response.json().get("response", "").strip()
#     summary = re.sub(r"<think>.*?</think>", "", summary, flags=re.DOTALL)
#     print("✅ Generated Summary:\n", summary)
#     return summary

# @app.route("/generate-preview", methods=["POST"])
# def generate_preview():
#     try:
#         print("🔥 /generate-preview called")
#         data = request.form.to_dict()
#         print("✅ Form data received:", data)

#         # Use base64 image if provided
#         photo_url = data.get("PHOTO_URL")

#         # If not base64, fall back to actual file upload
#         if not photo_url:
#             photo = request.files.get("photo")
#             if photo:
#                 filename = secure_filename(photo.filename)
#                 unique_filename = f"{uuid.uuid4().hex}_{filename}"
#                 save_path = os.path.join(UPLOAD_FOLDER, unique_filename)
#                 photo.save(save_path)
#                 photo_url = f"/uploads/{unique_filename}"
#             else:
#                 photo_url = "https://via.placeholder.com/120"
#         data["PHOTO_URL"] = photo_url

#         summary = generate_summary(data)
#         data["SUMMARY"] = summary

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": Markup(data.get("experience", "")),
#             "EDUCATION": Markup(data.get("education", "")),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": Markup(data.get("skills", "")),
#             "PROJECTS": Markup(data.get("projects", "")),
#             "AWARDS": Markup(data.get("awards", "")),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": summary,
#             "PHOTO_URL": photo_url
#         }

#         html_out = html_template.render(**context)
#         return jsonify({"html": html_out})

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/download-pdf", methods=["POST"])
# def download_pdf():
#     try:
#         data = request.form.to_dict()
#         summary = generate_summary(data)
#         data["SUMMARY"] = summary

#         photo_url = data.get("PHOTO_URL")
#         if not photo_url:
#             photo = request.files.get("photo")
#             if photo:
#                 filename = secure_filename(photo.filename)
#                 unique_filename = f"{uuid.uuid4().hex}_{filename}"
#                 save_path = os.path.join(UPLOAD_FOLDER, unique_filename)
#                 photo.save(save_path)
#                 photo_url = f"file://{os.path.abspath(save_path)}"
#             else:
#                 photo_url = "https://via.placeholder.com/120"
#         data["PHOTO_URL"] = photo_url

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": Markup(data.get("experience", "")),
#             "EDUCATION": Markup(data.get("education", "")),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": Markup(data.get("skills", "")),
#             "PROJECTS": Markup(data.get("projects", "")),
#             "AWARDS": Markup(data.get("awards", "")),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": summary,
#             "PHOTO_URL": photo_url
#         }

#         final_html = html_template.render(**context)
#         pdf_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.pdf")
#         HTML(string=final_html, base_url=os.getcwd()).write_pdf(pdf_path)

#         return send_file(pdf_path, as_attachment=True, download_name="resume.pdf")

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/uploads/<filename>")
# def uploaded_file(filename):
#     return send_from_directory(UPLOAD_FOLDER, filename)

# if __name__ == "__main__":
#     app.run(debug=True)


# # ✅ app.py – Backend for CV Generator with HTML Templating + Ollama Summary

# from flask import Flask, request, jsonify, send_file, render_template_string
# from flask_cors import CORS
# import os, uuid, re, requests, traceback
# from jinja2 import Template
# from markupsafe import Markup
# from weasyprint import HTML

# app = Flask(__name__)
# CORS(app)

# UPLOAD_FOLDER = "uploads"
# TEMPLATE_PATH = "preview.html"
# PROMPT_PATH = "prompt_template.txt"

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# # Load HTML template on startup
# with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
#     html_template = Template(f.read())

# # Load prompt template
# with open(PROMPT_PATH, "r") as f:
#     prompt_template = f.read()

# def generate_summary(data):
#     prompt = prompt_template.format(**data)
#     print("🧠 Prompt sent to Ollama:\n", prompt)  # 🔍 log full prompt
#     response = requests.post("http://localhost:11434/api/generate", json={
#         "model": "deepseek-r1:1.5b",
#         "prompt": prompt,
#         "stream": False
#     })

#     if response.status_code != 200:
#         raise Exception(f"Ollama error: {response.status_code}: {response.text}")

#     summary = response.json().get("response", "").strip()
#     summary = re.sub(r"<think>.*?</think>", "", summary, flags=re.DOTALL)
#     print("✅ Generated Summary:\n", summary)  # 🔍 log model output
#     return summary

# @app.route("/generate-preview", methods=["POST"])
# def generate_preview():
#     try:
#         print("🔥 /generate-preview called")
#         data = request.form.to_dict()
#         print("✅ Form data received:", data)
#         photo = request.files.get("photo")

#         if photo:
#             print("✅ Photo uploaded:", photo.filename)
#             photo_ext = photo.filename.split(".")[-1]
#             photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
#             photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
#             photo.save(photo_path)
#             data["PHOTO_URL"] = f"/uploads/{photo_filename}"
#         else:
#             data["PHOTO_URL"] = "https://via.placeholder.com/120"

#         summary = generate_summary(data)
#         data["SUMMARY"] = summary

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": data.get("experience", ""),
#             "EDUCATION": data.get("education", ""),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": data.get("skills", ""),
#             "PROJECTS": data.get("projects", ""),
#             "AWARDS": data.get("awards", ""),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": summary,
#             "PHOTO_URL": data["PHOTO_URL"]
#         }

#         for key in ["SKILLS", "PROJECTS", "AWARDS", "EXPERIENCE", "EDUCATION"]:
#             context[key] = Markup(context[key])

#         html_out = html_template.render(**context)
#         return jsonify({"html": html_out})

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/download-pdf", methods=["POST"])
# def download_pdf():
#     try:
#         data = request.form.to_dict()
#         summary = generate_summary(data)
#         data["SUMMARY"] = summary

#         photo = request.files.get("photo")
#         if photo:
#             photo_ext = photo.filename.split(".")[-1]
#             photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
#             photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
#             photo.save(photo_path)
#             data["PHOTO_URL"] = f"file://{os.path.abspath(photo_path)}"
#         else:
#             data["PHOTO_URL"] = "https://via.placeholder.com/120"

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": data.get("experience", ""),
#             "EDUCATION": data.get("education", ""),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": data.get("skills", ""),
#             "PROJECTS": data.get("projects", ""),
#             "AWARDS": data.get("awards", ""),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": summary,
#             "PHOTO_URL": data["PHOTO_URL"]
#         }

#         for key in ["SKILLS", "PROJECTS", "AWARDS", "EXPERIENCE", "EDUCATION"]:
#             context[key] = Markup(context[key])

#         final_html = html_template.render(**context)
#         pdf_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.pdf")
#         HTML(string=final_html, base_url=os.getcwd()).write_pdf(pdf_path)

#         return send_file(pdf_path, as_attachment=True, download_name="resume.pdf")

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/uploads/<filename>")
# def uploaded_file(filename):
#     return send_file(os.path.join(UPLOAD_FOLDER, filename))

# if __name__ == "__main__":
#     app.run(debug=True)





# # ✅ app.py – Backend for CV Generator with HTML Templating + Ollama Summary

# from flask import Flask, request, jsonify, send_file, render_template_string
# from flask_cors import CORS
# import os, uuid, re, requests, traceback
# from jinja2 import Template
# from weasyprint import HTML

# app = Flask(__name__)
# CORS(app)

# UPLOAD_FOLDER = "uploads"
# TEMPLATE_PATH = "preview.html"
# PROMPT_PATH = "prompt_template.txt"

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# # Load HTML template on startup
# with open(TEMPLATE_PATH, "r", encoding="utf-8") as f:
#     html_template = Template(f.read())

# # Load prompt template
# with open(PROMPT_PATH, "r") as f:
#     prompt_template = f.read()

# def generate_summary(data):
#     prompt = prompt_template.format(**data)
#     print("🧠 Prompt sent to Ollama:\n", prompt)  # 🔍 log full prompt
#     response = requests.post("http://localhost:11434/api/generate", json={
#         "model": "deepseek-r1:1.5b",
#         "prompt": prompt,
#         "stream": False
#     })

#     if response.status_code != 200:
#         raise Exception(f"Ollama error: {response.status_code}: {response.text}")

#     summary = response.json().get("response", "").strip()
#     summary = re.sub(r"<think>.*?</think>", "", summary, flags=re.DOTALL)
#     print("✅ Generated Summary:\n", summary)  # 🔍 log model output
#     return summary

# @app.route("/generate-preview", methods=["POST"])
# def generate_preview():
#     try:
#         print("🔥 /generate-preview called")  # 🔍 entry point
#         data = request.form.to_dict()
#         print("✅ Form data received:", data)  # 🔍 log form data
#         photo = request.files.get("photo")

#         if photo:
#             print("✅ Photo uploaded:", photo.filename)  # 🔍 log uploaded photo
#             photo_ext = photo.filename.split(".")[-1]
#             photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
#             photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
#             photo.save(photo_path)
#             data["PHOTO_URL"] = f"/uploads/{photo_filename}"
#         else:
#             data["PHOTO_URL"] = "https://via.placeholder.com/120"  # fallback image

#         # Call Ollama only for summary
#         summary = generate_summary(data)
#         data["SUMMARY"] = summary

#         # Map placeholders in template
#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": data.get("experience", ""),
#             "EDUCATION": data.get("education", ""),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": data.get("skills", ""),
#             "PROJECTS": data.get("projects", ""),
#             "AWARDS": data.get("awards", ""),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": summary,
#             "PHOTO_URL": data["PHOTO_URL"]
#         }

#         html_out = html_template.render(**context)
#         return jsonify({"html": html_out})

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/download-pdf", methods=["POST"])
# def download_pdf():
#     try:
#         data = request.form.to_dict()
#         summary = generate_summary(data)
#         data["SUMMARY"] = summary

#         # Save uploaded image
#         photo = request.files.get("photo")
#         if photo:
#             photo_ext = photo.filename.split(".")[-1]
#             photo_filename = f"{uuid.uuid4().hex}.{photo_ext}"
#             photo_path = os.path.join(UPLOAD_FOLDER, photo_filename)
#             photo.save(photo_path)
#             data["PHOTO_URL"] = f"file://{os.path.abspath(photo_path)}"
#         else:
#             data["PHOTO_URL"] = "https://via.placeholder.com/120"

#         context = {
#             "NAME": data.get("name", ""),
#             "ADDRESS": data.get("address", ""),
#             "EMAIL": data.get("email", ""),
#             "NUMBER": data.get("phone", ""),
#             "ROLE": data.get("role", ""),
#             "EXPERIENCE": data.get("experience", ""),
#             "EDUCATION": data.get("education", ""),
#             "PROFESSIONAL": data.get("prof_education", ""),
#             "SKILLS": data.get("skills", ""),
#             "PROJECTS": data.get("projects", ""),
#             "AWARDS": data.get("awards", ""),
#             "START": data.get("start", ""),
#             "END": data.get("end", ""),
#             "SUMMARY": summary,
#             "PHOTO_URL": data["PHOTO_URL"]
#         }

#         final_html = html_template.render(**context)
#         pdf_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.pdf")
#         HTML(string=final_html, base_url=os.getcwd()).write_pdf(pdf_path)

#         return send_file(pdf_path, as_attachment=True, download_name="resume.pdf")

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route("/uploads/<filename>")
# def uploaded_file(filename):
#     return send_file(os.path.join(UPLOAD_FOLDER, filename))

# if __name__ == "__main__":
#     app.run(debug=True)





# # 📁 app.py

# from flask import Flask, request, jsonify, send_file, send_from_directory
# from flask_cors import CORS
# import os, uuid, re, traceback, requests
# from docx import Document
# from docx.shared import Inches

# app = Flask(__name__)
# CORS(app)

# UPLOAD_FOLDER = "uploads"
# OUTPUT_FOLDER = "generated"
# TEMPLATE_PATH = "cv-template-filled (3).docx"  # Your Word template

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# # Load the prompt template (for Ollama)
# with open("prompt_template.txt", "r") as f:
#     template = f.read()

# # 🔧 Replace text placeholders
# def replace_placeholders(doc, replacements):
#     for para in doc.paragraphs:
#         for run in para.runs:
#             for key, val in replacements.items():
#                 if key in run.text:
#                     run.text = run.text.replace(key, val)

# # 🧠 Generate LLM output using Ollama
# def prompting_cv(data):
#     prompt_text = template.format(**data)

#     response = requests.post(
#         "http://localhost:11434/api/generate",
#         json={
#             "model": "deepseek-r1:1.5b",
#             "prompt": prompt_text,
#             "stream": False
#         }
#     )

#     if response.status_code != 200:
#         raise Exception(f"Ollama error {response.status_code}: {response.text}")

#     raw_response = response.json()["response"]
#     clean_response = re.sub(r"<think>.*?</think>", "", raw_response, flags=re.DOTALL).strip()

#     # Extract sections from LLM response
#     sections = {}
#     for section in ["SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS", "AWARDS", "PROFESSIONAL_QUALIFICATIONS"]:
#         match = re.search(fr"{section}:(.*?)(?=\n[A-Z_]+:|\Z)", clean_response, re.DOTALL)
#         if match:
#             sections[section] = match.group(1).strip()
#     return sections

# # 📄 Route to generate CV (.docx)
# @app.route('/generate-cv', methods=['POST'])
# def generate_cv():
#     try:
#         data = request.form.to_dict()
#         photo = request.files.get("photo")
#         image_path = None

#         if photo:
#             image_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.png")
#             photo.save(image_path)

#         result = prompting_cv(data)

#         filled = {
#             "{{NAME}}": data.get("name", ""),
#             "{{ADDRESS}}": data.get("address", ""),
#             "{{EMAIL}}": data.get("email", ""),
#             "{{NUMBER}}": data.get("phone", ""),
#             "{{ROLE}}": data.get("role", ""),
#             "{{SUMMARY}}": result.get("SUMMARY", ""),
#             "{{EXPERIENCE}}": result.get("EXPERIENCE", ""),
#             "{{EDUCATION}}": result.get("EDUCATION", ""),
#             "{{PROFESSIONAL}}": result.get("PROFESSIONAL_QUALIFICATIONS", ""),
#             "{{SKILLS}}": result.get("SKILLS", ""),
#             "{{PROJECTS}}": result.get("PROJECTS", ""),
#             "{{AWARDS}}": result.get("AWARDS", "")
#         }

#         doc = Document(TEMPLATE_PATH)

#         # 🔁 Replace all text placeholders
#         replace_placeholders(doc, filled)

#         # 📸 Replace {{PHOTO}} with actual image
#         if image_path:
#             for para in doc.paragraphs:
#                 if "{{PHOTO}}" in para.text:
#                     para.clear()  # Remove the placeholder text
#                     run = para.add_run()
#                     run.add_picture(image_path, width=Inches(1.5))
#                     break

#         output_path = os.path.join(OUTPUT_FOLDER, f"{uuid.uuid4().hex}.docx")
#         doc.save(output_path)

#         return send_file(output_path, as_attachment=False, download_name="Generated_CV.docx")

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# # 📄 Route to preview PDFs (optional)
# @app.route('/preview/<filename>')
# def preview(filename):
#     return send_from_directory(OUTPUT_FOLDER, filename)

# if __name__ == "__main__":
#     app.run(debug=True)












# # ✅ app.py (Backend: Flask + Ollama + docx2pdf + live preview)

# from flask import Flask, request, jsonify, send_from_directory
# from flask_cors import CORS
# import traceback
# import requests
# import re
# import os
# import uuid
# import pythoncom
# import win32com.client
# from docx import Document
# from docx.shared import Inches
# from docx2pdf import convert

# app = Flask(__name__)
# CORS(app)

# UPLOAD_FOLDER = "uploads"
# OUTPUT_FOLDER = "generated"
# TEMPLATE_PATH = "cv_template.docx"  # Ensure this template exists

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# with open("prompt_template.txt", "r") as f:
#     template = f.read()

# def prompting_cv(data):
#     prompt_text = template.format(**data)
#     response = requests.post(
#         "http://localhost:11434/api/generate",
#         json={"model": "deepseek-r1:1.5b", "prompt": prompt_text, "stream": False}
#     )
#     if response.status_code != 200:
#         raise Exception(f"Ollama error {response.status_code}: {response.text}")

#     raw_response = response.json()["response"]
#     clean_response = re.sub(r"<think>.*?</think>", "", raw_response, flags=re.DOTALL).strip()

#     sections = {}
#     for section in ["SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS", "AWARDS", "PROFESSIONAL_QUALIFICATIONS"]:
#         match = re.search(fr"{section}:(.*?)(?=\n[A-Z_]+:|\Z)", clean_response, re.DOTALL)
#         if match:
#             sections[section] = match.group(1).strip()
#     return sections

# @app.route('/generate-cv', methods=['POST'])
# def generate_cv():
#     try:
#         data = request.form.to_dict()
#         photo = request.files.get('photo')

#         image_path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4().hex}.png")
#         if photo:
#             photo.save(image_path)

#         sections = prompting_cv(data)

#         filled = {
#             'NAME': data.get('name'),
#             'EMAIL': data.get('email'),
#             'NUMBER': data.get('phone'),
#             'ADDRESS': data.get('address'),
#             'ROLE': data.get('role'),
#             'SUMMARY': sections.get('SUMMARY', ''),
#             'EXPERIENCE': sections.get('EXPERIENCE', ''),
#             'EDUCATION': sections.get('EDUCATION', ''),
#             'PROFESSIONAL': sections.get('PROFESSIONAL_QUALIFICATIONS', ''),
#             'SKILLS': sections.get('SKILLS', ''),
#             'PROJECTS': sections.get('PROJECTS', ''),
#             'AWARDS': sections.get('AWARDS', '')
#         }

#         doc = Document("cv_template_filled (2).docx")
#         for p in doc.paragraphs:
#             for key, value in filled.items():
#                 if f"{{{{{key}}}}}" in p.text:
#                     p.text = p.text.replace(f"{{{{{key}}}}}", value)

#         if photo:
#             doc.paragraphs[0].insert_paragraph_before().add_run().add_picture(image_path, width=Inches(1.5))

#         docx_filename = f"{uuid.uuid4().hex}.docx"
#         docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)
#         doc.save(docx_path)

#         pythoncom.CoInitialize()
#         pdf_path = docx_path.replace(".docx", ".pdf")
#         convert(docx_path, pdf_path)

#         return jsonify({
#             "docx_url": f"http://localhost:5000/preview/{os.path.basename(docx_path)}",
#             "pdf_url": f"http://localhost:5000/preview/{os.path.basename(pdf_path)}"
#         })

#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500

# @app.route('/preview/<path:filename>')
# def preview_file(filename):
#     return send_from_directory(OUTPUT_FOLDER, filename)

# if __name__ == '__main__':
#     app.run(debug=True)






# ✅ Gemini-based CV Generator Backend (Flask)

# from flask import Flask, request, send_file
# from flask_cors import CORS
# from docx import Document
# from docx.shared import Inches
# from io import BytesIO
# import tempfile, os, re
# from docx2pdf import convert
# import google.generativeai as genai

# # Initialize Flask app
# app = Flask(__name__)
# CORS(app, resources={r"/*": {"origins": "*"}})  # for development only

# # 🔐 Configure Gemini API Key from environment variable
# from dotenv import load_dotenv
# load_dotenv()
# genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# # Load prompt template
# with open("prompt_template.txt", "r") as f:
#     template = f.read()

# def prompting_cv(data):
#     prompt_text = template.format(**data)

#     print(">>> Prompt sent to Ollama:\n", prompt_text)

#     response = requests.post(
#         "http://localhost:11434/api/generate",
#         json={
#             "model": "deepseek-r1:1.5b",
#             "prompt": prompt_text,
#             "stream": False
#         }
#     )

#     if response.status_code != 200:
#         raise Exception(f"Ollama error {response.status_code}: {response.text}")

#     raw_response = response.json()["response"]
#     clean_response = re.sub(r"<think>.*?</think>", "", raw_response, flags=re.DOTALL).strip()

#     # Parse sections
#     sections = {}
#     for section in ["SUMMARY", "EXPERIENCE", "EDUCATION", "PROFESSIONAL QUALIFICATIONS", "SKILLS", "PROJECTS", "AWARDS"]:
#         match = re.search(rf"{section}:\s*(.*?)(?=\n[A-Z ]+?:|\Z)", clean_response, re.DOTALL)
#         key = section.replace(" ", "_")
#         sections[key] = match.group(1).strip() if match else ""

#     return sections
#     # prompt_text = template.format(**data)
#     # print(">>> Prompt sent to Gemini:\n", prompt_text)
#     # model = genai.GenerativeModel("gemini-1.5-pro")
#     # response = model.generate_content(prompt_text)
#     # return response.text.strip()

# def replace_first_image(doc, photo_path):
#     for shape in doc.inline_shapes:
#         rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
#         new_rid = doc.part.relate_to(photo_path,
#             "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
#             is_external=False)
#         shape._inline.graphic.graphicData.pic.blipFill.blip.embed = new_rid
#         break

# def fill_template(cv_text, doc_path, photo_file):
#     doc = Document(doc_path)
#     for para in doc.paragraphs:
#         for run in para.runs:
#             run.text = run.text.replace("{{CV_BODY}}", cv_text)
#     if photo_file:
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
#             photo_file.save(tmp.name)
#             replace_first_image(doc, tmp.name)
#             os.unlink(tmp.name)
#     return doc

# @app.route("/generate-cv", methods=["POST"])
# def generate_docx():
#     form = request.form
#     photo = request.files.get("photo")
#     data = {
#         "name": form.get("name", ""),
#         "address": form.get("address", ""),
#         "email": form.get("email", ""),
#         "phone": form.get("phone", ""),
#         "summary": form.get("summary", ""),
#         "role": form.get("role", ""),
#         "experience": form.get("experience", ""),
#         "education": form.get("education", ""),
#         "professional": form.get("professional", ""),
#         "skills": form.get("skills", ""),
#         "projects": form.get("projects", ""),
#         "awards": form.get("awards", ""),
#         "duration": form.get("duration", "")
#     }
#     result = prompting_cv(data)
#     doc = fill_template(result, "cv-template-filled_updated.docx", photo)
#     buf = BytesIO()
#     doc.save(buf)
#     buf.seek(0)
#     return send_file(buf, as_attachment=True, download_name="Generated_CV.docx")

# @app.route("/generate-pdf", methods=["POST"])
# def generate_pdf():
#     form = request.form
#     photo = request.files.get("photo")
#     data = {
#         "name": form.get("name", ""),
#         "address": form.get("address", ""),
#         "email": form.get("email", ""),
#         "phone": form.get("phone", ""),
#         "summary": form.get("summary", ""),
#         "role": form.get("role", ""),
#         "experience": form.get("experience", ""),
#         "education": form.get("education", ""),
#         "professional": form.get("professional", ""),
#         "skills": form.get("skills", ""),
#         "projects": form.get("projects", ""),
#         "awards": form.get("awards", ""),
#         "duration": form.get("duration", "")
#     }
#     result = prompting_cv(data)
#     doc = fill_template(result, "cv-template-filled.docx", photo)
#     with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_file:
#         doc.save(docx_file.name)
#         pdf_path = docx_file.name.replace(".docx", ".pdf")
#         convert(docx_file.name, pdf_path)
#         with open(pdf_path, "rb") as pdf:
#             pdf_bytes = BytesIO(pdf.read())
#         os.remove(docx_file.name)
#         os.remove(pdf_path)
#         pdf_bytes.seek(0)
#         return send_file(pdf_bytes, as_attachment=True, download_name="Generated_CV.pdf")

# if __name__ == "__main__":
#     app.run(debug=True)








# # backend/app.py

# from flask import Flask, request, send_file
# from flask_cors import CORS
# from docx import Document
# from docx.shared import Inches
# from io import BytesIO
# import tempfile, os, re, requests
# from docx2pdf import convert

# app = Flask(__name__)
# CORS(app)

# # Load prompt template
# with open("prompt_template.txt", "r") as f:
#     template = f.read()

# def prompting_cv(data):
#     prompt_text = template.format(**data)
#     response = requests.post("http://localhost:11434/api/generate", json={
#         "model": "deepseek-r1:1.5b",
#         "prompt": prompt_text,
#         "stream": False
#     })
#     raw = response.json()["response"]
#     return re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()

# def replace_first_image(doc, photo_path):
#     for shape in doc.inline_shapes:
#         rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
#         new_rid = doc.part.relate_to(photo_path,
#             "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
#             is_external=False)
#         shape._inline.graphic.graphicData.pic.blipFill.blip.embed = new_rid
#         break

# def fill_template(content, doc_path, photo_file):
#     doc = Document(doc_path)
#     for para in doc.paragraphs:
#         for run in para.runs:
#             run.text = run.text.replace("{{CV_BODY}}", content)
#     if photo_file:
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
#             photo_file.save(tmp.name)
#             replace_first_image(doc, tmp.name)
#             os.unlink(tmp.name)
#     return doc

# @app.route("/generate-cv", methods=["POST"])
# def generate_docx():
#     form = request.form
#     photo = request.files.get("photo")
#     data = {k: form.get(k, "") for k in ["name", "address", "email", "phone",  "summary","role", "experience" ,"education","prof_education", "skills", "projects", "awards"]}
#     result = prompting_cv(data)
#     doc = fill_template(result, "cv-template-filled.docx", photo)
#     buf = BytesIO()
#     doc.save(buf)
#     buf.seek(0)
#     return send_file(buf, as_attachment=True, download_name="Generated_CV.docx")

# @app.route("/generate-pdf", methods=["POST"])
# def generate_pdf():
#     form = request.form
#     photo = request.files.get("photo")
#     data = {k: form.get(k, "") for k in ["name", "address", "email", "phone",  "summary","role", "experience" "education","prof_education" "skills", "projects", "awards"]}
#     result = prompting_cv(data)
#     doc = fill_template(result, "cv-template-filled.docx", photo)
#     with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_file:
#         doc.save(docx_file.name)
#         pdf_path = docx_file.name.replace(".docx", ".pdf")
#         convert(docx_file.name, pdf_path)
#         with open(pdf_path, "rb") as pdf:
#             pdf_bytes = BytesIO(pdf.read())
#         os.remove(docx_file.name)
#         os.remove(pdf_path)
#         pdf_bytes.seek(0)
#         return send_file(pdf_bytes, as_attachment=True, download_name="Generated_CV.pdf")

# if __name__ == "__main__":
#     app.run(debug=True)
















# from flask import Flask, request, jsonify, send_file
# from flask_cors import CORS
# import traceback
# import requests
# import json
# import re
# from io import BytesIO
# from reportlab.pdfgen import canvas

# app = Flask(__name__)
# CORS(app)  

# # Load the prompt template from a file
# with open("prompt_template.txt", "r") as f:
#     template = f.read()

# def prompting_cv(data):
#     """Generate resume using Ollama deepseek-r1:1.5b locally."""
#     # Fill in the prompt template with user data
#     prompt_text = template.format(**data)
        

#     # Call Ollama local API
#     print(">>> Prompt to Ollama:\n", prompt_text)  

#     response = requests.post(
#         "http://localhost:11434/api/generate",
#         json={
#             "model": "deepseek-r1:1.5b",
#             "prompt": prompt_text,
#             "stream": False
#         }
#     )

#     print(">>> Ollama response status:", response.status_code)  
#     print(">>> Ollama response body:", response.text[:500])     

#     if response.status_code != 200:
#         raise Exception(f"Ollama error {response.status_code}: {response.text}")
#      # Extract and clean response
#     raw_response = response.json()["response"]

#     # 🧽 Remove <think>...</think> section (if any)
#     clean_response = re.sub(r"<think>.*?</think>", "", raw_response, flags=re.DOTALL).strip()

#     return clean_response


# @app.route('/generate-cv', methods=['POST'])
# def generate_cv():
#     data = request.json

#     required_fields = [
#         "name", "role", "email", "phone", "address",
#         "summary", "education", "experience",
#         "skills", "projects", "awards"
#     ]

#     for field in required_fields:
#         if field not in data:
#             return jsonify({"error": f"Missing field: {field}"}), 400

#     try:
#         result = prompting_cv(data)
        

#     # Create PDF
#         buffer = BytesIO()
#         p = canvas.Canvas(buffer)
#         y = 800
#         for line in result.split('\n'):

#             p.drawString(50, y, line)
#             y -= 20
#             if y < 40:

#                 p.showPage()
#                 y = 800
#         p.save()
#         buffer.seek(0)

#         return send_file(buffer, as_attachment=True, download_name="resume.pdf", mimetype='application/pdf')
        
#     except Exception as e:
#         traceback.print_exc()
#         return jsonify({"error": str(e)}), 500




# if __name__ == '__main__':
#     app.run(debug=True)






